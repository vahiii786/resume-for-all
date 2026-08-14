
Claude finished the response

import streamlit as st import re from collections import Counter import math import urllib.parse from PIL import Image # Multi-level PDF & File Extraction Libraries fitz_available = False try: import fitz # PyMuPDF fitz_available = True except ImportError: fitz_available =

pasted

idhi naa resume app ki sambandhinchina code so indhulo madyalo telugucontent undhi adhi unwanted ayite remove cheyyu importent ayite naa app lo aa telugu content show avvakunda hide cheyu


Identified Telugu text and strategized localization approach efficiently
Identified Telugu text and strategized localization approach efficiently
ఈ కోడ్‌లో Telugu content 4 చోట్ల ఉంది — అన్నీ UI లో కనిపించే helper/description టెక్స్ట్‌లే (functionality కి ఏమీ సంబంధం లేదు), కాబట్టి వాటిని English కి మార్చేసాను. దీనివల్ల app లో ఇక Telugu కనిపించదు, కానీ ఆ description సమాచారం (ఏం చెయ్యాలి అనే గైడెన్స్) పోకుండా అలాగే ఉంటుంది.

Changed lines:

Tab 2 description → English
Tab 6 description → English
Tab 7 description → English
One comment (# Mandatory కాకుండా Optional JD Toggle) → English
Rest of the code (logic, variable names, functions) — అస్సలు మార్చలేదు, safe గా ఉంచాను.


Created a file, read a file
Created a file, read a file
ఇక app run చేస్తే Telugu text ఎక్కడా కనిపించదు — logic, features అన్నీ యధావిధిగానే పనిచేస్తాయి.


Resume app
PY 
[Resume-for-All — Explainable Resume Analysis](https://resume-insight-lab.preview.emergentagent.com/)
ee link asala ento chudu ready chey andhulo data enti


Fetched: Resume-for-All — Explainable Resume Analysis
Fetched: Resume-for-All — Explainable Resume Analysis
ఆ link "Resume-for-All" అనే ఒక resume analysis tool. దీని గురించి ఏం ఉందో చూసాను:

ఏం చేస్తుంది:

మీ resume (PDF/DOCX/TXT, max 8MB) upload చేయాలి లేదా text paste చేయాలి
Job Description (JD) optional గా paste చెయ్యొచ్చు — match score కోసం
"Analyze Resume" బటన్ నొక్కితే resume ని analyze చేస్తుంది
Features (వాళ్ళు చెప్పుకున్నవి):

Explainable score — score కి ప్రతి పాయింట్‌కి verbatim evidence తో reason చూపిస్తుంది
No fabrication — AI మీరు ఇవ్వని metrics ని invent చెయ్యదు (అంటే fake numbers add చేయదు)
Export PDF/DOCX — improved resume ని ఒక్క క్లిక్‌లో download చేసుకోవచ్చు
Tech: GPT-4o mini వాడుతున్నారని చెప్పుకుంది, Free · Anonymous · Honest అని tagline.

"అందులో data ఏంటి" అంటే — ఇది కేవలం ఒక landing page/UI మాత్రమే (static homepage), actual analysis logic/backend ఏం చేస్తుందో ఈ page లో చూపించలేదు. బహుశా ఇది మీ resume app లాగే ఒక similar streamlit/web tool అయ్యుండొచ్చు, కానీ backend logic (scoring algorithm, prompt) ఏంటో ఈ page నుండి తెలియదు — అది వాళ్ళ server-side code లో ఉంటుంది, ఇక్కడ కనిపించదు.

మీరు మీ app లో ఇలాంటి "Explainable score with verbatim evidence" feature add చేయాలనుకుంటున్నారా? అలా అయితే, ఇప్పుడున్న cosine-similarity match score కి, ఏ ఏ keyword/sentence match అయ్యిందో reason గా చూపించే logic add చేయగలను.

haha add cheste baguntadha and features kuda

అవును, add చేస్తే బాగుంటుంది — ఎందుకంటే ఇప్పుడు మీ app లో match score కేవలం ఒక సంఖ్య మాత్రమే చూపిస్తుంది (cosine similarity %), కానీ "ఎందుకు ఈ score వచ్చింది" అని user కి కనిపించదు. "Explainable" feature add చేస్తే trust పెరుగుతుంది, resume ఎక్కడ improve చెయ్యాలో clear గా తెలుస్తుంది.

కొన్ని బాగుండే feature ideas:

Evidence-based match — ప్రతి matched keyword కి, resume లో ఏ line/bullet లో అది కనిపించిందో చూపించడం (just score కాకుండా "why")
Missing keyword → exact fix suggestion — మీ Tab 4 లో ఇది కొంతవరకు ఉంది, దాన్ని improve చేయొచ్చు
No-fabrication rewrite rule — Tab 3 (Bullet Rewriter) లో స్టాటిక్ templates బదులు, resume లో ఉన్న actual facts/numbers మాత్రమే వాడి rewrite చేయడం (fake metrics add చేయకుండా)
Section-wise score breakdown — overall score కాకుండా, Skills match %, Experience match %, Keywords match % విడివిడిగా చూపించడం
Export improved resume as PDF/DOCX — ప్రస్తుతం మీ app HTML మాత్రమే download చేస్తుంది, real PDF/DOCX export add చేస్తే బాగుంటుంది
ఇవన్నీ మీ ప్రస్తుత app architecture (cosine similarity + keyword matching) మీదే build చేయొచ్చు, పెద్ద rewrite అవసరం లేదు.

ok add chesi pettu with new features


Architected comprehensive resume analyzer with section scoring, evidence matching, and export functionality
Architected comprehensive resume analyzer with section scoring, evidence matching, and export functionality
బాగుంటుంది! కింద features add చేసి పూర్తి app ready చేస్తున్నాను:

Tab 1 → Section-wise score breakdown (Skills/Experience/Projects/Education % విడివిడిగా) + Evidence-based matched keywords (ఏ keyword resume లో ఏ line లో దొరికిందో చూపిస్తుంది)
Tab 2 → real DOCX & PDF export buttons (ఇప్పటివరకు HTML మాత్రమే ఉండేది)
Tab 3 → No-fabrication bullet rewriter — మీ original bullet లో number ఉంటేనే దాన్ని వాడుతుంది, లేకపోతే fake % invent చేయదు (transparent గా చెప్తుంది కూడా)

Created 2 files, ran 2 commands, viewed a file
Created 2 files, ran 2 commands, viewed a file
Ready — syntax check pass ఐంది. కొత్తగా ఏం add చేసానో:

Tab 1 (Analyzer) — JD ఉన్నప్పుడు:

🧩 Section-wise breakdown: Skills / Experience / Projects / Education match % విడివిడిగా
🔎 Evidence list: matched keyword + అది resume లో ఏ exact line లో దొరికిందో
Tab 2 (Builder) — Download row లో ఇప్పుడు 3 బటన్లు: HTML (పాతది) + DOCX + PDF (కొత్తవి, python-docx & fpdf2 వాడి)

Tab 3 (Bullet Rewriter) — original line లో number ఉంటే దాన్నే వాడుతుంది; లేకపోతే fake % ఏమీ కలపదు, పైగా "no number found so none was invented" అని చెప్తుంది

Run చేయడానికి కావాల్సిన కొత్త packages:

pip install python-docx fpdf2
(ఇవి లేకపోయినా app crash అవ్వదు — ఆ export బటన్లు మాత్రం disabled గా చూపిస్తుంది, install చెయ్యమని hint ఇస్తుంది.)


Resume app
PY 
nenu install chesinakuda pdf dhi alne vastundhi


Diagnosed persistent PDF functionality issue despite installation
Diagnosed persistent PDF functionality issue despite installation
ఇది సాధారణంగా 4 కారణాల్లో ఒకటి వల్ల జరుగుతుంది — చెక్ చేద్దాం:

1. App restart అవ్వలేదు
pip install చేసినా, Streamlit app ఇంకా పాత process లోనే running ఉంటే కొత్త package కనపడదు. కేవలం page refresh సరిపోదు — terminal లో app ని పూర్తిగా ఆపి (Ctrl+C) మళ్ళీ streamlit run resume_app.py అని start చేయాలి.

2. వేరే Python environment లో install అయ్యింది
మీరు pip install run చేసిన terminal, మరియు streamlit app run అయ్యే terminal/environment వేరైతే (virtualenv, conda env, లేదా online IDE లో వేరే kernel), package కనపడదు. చెక్ చెయ్యండి:

bash
which python
which streamlit
pip show fpdf2
మూడూ ఒకే environment చూపిస్తున్నాయో లేదో confirm చేసుకోండి.

3. Deployed/hosted app అయితే (Streamlit Cloud, Emergent, etc.)
మీరు local terminal లో pip install చేస్తే అది మీ machine లో మాత్రమే install అవుతుంది — cloud/hosted deployment మీద install అవ్వదు. అలాంటప్పుడు requirements.txt ఫైల్‌లో ఇలా add చేసి, redeploy/reboot చేయాలి:

python-docx
fpdf2
4. fpdf vs fpdf2 package name conflict
పాత fpdf package కూడా install అయి ఉంటే conflict రావొచ్చు. క్లీన్‌గా చెయ్యడానికి:

bash
pip uninstall fpdf fpdf2 -y
pip install fpdf2
మీరు ఏ విధంగా app run చేస్తున్నారు (local terminal లో streamlit run, లేదా ఏదైనా hosted platform మీద)? అది చెప్తే సరైన కారణం ఏంటో exact గా చెప్పగలను.

ayina alane vastundhi


Analyzed query and formulated response
Analyzed query and formulated response
సరే, exact గా ఎక్కడ fail అవుతుందో తెలుసుకోడానికి, code లోనే error message చూపించేలా మార్చేద్దాం — అప్పుడు "install లేదు" అని generic గా కాకుండా, actual Python error ఏంటో మీకే కనిపిస్తుంది.


Edited 4 files, ran a command, read a file
Edited 4 files, ran a command, read a file
ఇప్పుడు కొత్త file లో diagnostics add చేసాను. ఇలా చెయ్యండి:

కొత్త file డౌన్‌లోడ్ చేసుకుని app run చేయండి
Tab 2 లో export బటన్లు ఇంకా disabled గా ఉంటే, కింద "🔧 Why is export disabled? (diagnostics)" అనే expander కనిపిస్తుంది — దాన్ని click చేయండి
అందులో actual Python error message, మరియు app ఏ Python executable వాడుతోందో చూపిస్తుంది
ఆ expander లో వచ్చిన error text ని ఇక్కడ paste చేయండి — అప్పుడు exact గా ఏం సమస్యో చెప్పగలను (ఉదా: package name mismatch, version conflict, wrong environment, etc.).


Resume app
PY 

Message collapsed

Analyzed PDF generation code and encoding challenges
Analyzed PDF generation code and encoding challenges
ఇది ఒక వేరే/సులభమైన PDF generation code — మీ resume text ని PDF గా export చేయడానికి fpdf library వాడి రాసినది. Explain చేస్తాను:

ఏం చేస్తుంది:

PDF అనే class ని FPDF నుండి inherit చేసి define చేస్తుంది
body() method — text ని PDF లో రాస్తుంది, కానీ ముందు దాన్ని latin-1 encoding కి convert చేస్తుంది (Telugu లాంటి Unicode characters ఉంటే error రాకుండా, వాటిని ? లాంటి replacement characters తో replace చేస్తుంది — encode("latin-1", "replace"))
pdf.output(dest="S") — PDF ని string/bytes గా return చేస్తుంది
st.download_button — ఆ PDF bytes ని download చేసుకోడానికి బటన్ చూపిస్తుంది
ముఖ్యమైన విషయం — ఇదే మీ సమస్యకి కారణం కావొచ్చు:

pdf.output(dest="S") అనేది పాత fpdf library యొక్క API. కానీ fpdf2 (కొత్త, maintained library) లో ఇది మారింది — dest parameter కొన్ని versions లో deprecated/removed అయ్యింది, మరియు కొత్త fpdf2 versions లో pdf.output() నేరుగా bytearray return చేస్తుంది (encode చేయాల్సిన అవసరం లేదు).

కాబట్టి మీరు fpdf2 install చేసినా, ఈ code పాత API (dest="S") వాడుతుంటే TypeError లేదా similar error వస్తుంది — అదే మీకు "install చేసినా అలానే వస్తుంది" అనే సమస్యకి root cause అయ్యుండొచ్చు.

Fix: రెండు versions తోనూ పనిచేసేలా safe గా రాయాలి:


Ran a command
Ran a command


You are out of free messages until 7:10 PM




Claude is AI and can make mistakes. Please double-check responses.
Resume app · PY
import streamlit as st
import re
import io
import sys
from collections import Counter
import math
import urllib.parse
from PIL import Image
 
# Multi-level PDF & File Extraction Libraries
fitz_available = False
try:
    import fitz  # PyMuPDF
    fitz_available = True
except ImportError:
    fitz_available = False
 
pdfplumber_available = False
try:
    import pdfplumber
    pdfplumber_available = True
except ImportError:
    pdfplumber_available = False
 
pypdf_available = False
try:
    import pypdf
    pypdf_available = True
except ImportError:
    pypdf_available = False
 
docx_available = False
docx_import_error = ""
try:
    import docx  # python-docx (used for both reading uploads and writing exports)
    docx_available = True
except Exception as e:
    docx_available = False
    docx_import_error = str(e)
 
ocr_available = False
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    ocr_available = True
except ImportError:
    ocr_available = False
 
fpdf_available = False
fpdf_import_error = ""
try:
    from fpdf import FPDF
    fpdf_available = True
except Exception as e:
    fpdf_available = False
    fpdf_import_error = str(e)
 
st.set_page_config(
    page_title="AI Resume Suite & Career Hub",
    page_icon="🚀",
    layout="wide"
)
 
st.title("🚀 Smart AI Resume Hub & All-in-One Career Suite")
 
# Global Session State
if "built_resume_text" not in st.session_state:
    st.session_state["built_resume_text"] = ""
 
# All 7 Tabs Preserved
tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
    "📊 Resume Analyzer & Jobs", 
    "✏️ Live Resume Builder",
    "✨ AI Bullet Rewriter",
    "🎯 Keyword Placement & Red Flags", 
    "💰 Salary Predictor", 
    "✉️ HR Cold Outreach", 
    "🎤 AI Mock Interview"
])
 
# 4-Layer Bulletproof Text Extractor
def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return ""
    
    file_type = uploaded_file.name.split('.')[-1].lower()
    text = ""
    
    try:
        uploaded_file.seek(0)
        file_bytes = uploaded_file.read()
        
        if file_type == "pdf":
            # Layer 1: PyMuPDF (Fastest for standard text PDFs)
            if fitz_available:
                try:
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    for page in doc:
                        extracted = page.get_text()
                        if extracted:
                            text += extracted + "\n"
                except Exception:
                    text = ""
 
            # Layer 2: pdfplumber fallback
            if not text.strip() and pdfplumber_available:
                try:
                    uploaded_file.seek(0)
                    with pdfplumber.open(uploaded_file) as pdf:
                        for page in pdf.pages:
                            extracted = page.extract_text()
                            if extracted:
                                text += extracted + "\n"
                except Exception:
                    text = ""
 
            # Layer 3: pypdf fallback
            if not text.strip() and pypdf_available:
                try:
                    uploaded_file.seek(0)
                    reader = pypdf.PdfReader(uploaded_file)
                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
                except Exception:
                    text = ""
 
            # Layer 4: OCR Engine for Scanned / Canva / Image PDFs
            if not text.strip() and ocr_available:
                try:
                    images = convert_from_bytes(file_bytes)
                    for img in images:
                        extracted = pytesseract.image_to_string(img)
                        if extracted:
                            text += extracted + "\n"
                except Exception:
                    text = ""
 
        elif file_type in ["docx", "doc"]:
            if docx_available:
                uploaded_file.seek(0)
                doc = docx.Document(uploaded_file)
                for para in doc.paragraphs:
                    if para.text:
                        text += para.text + "\n"
 
        elif file_type in ["txt", "md"]:
            text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
            
    except Exception:
        pass
        
    return text.strip()
 
def text_to_vector(text):
    words = re.findall(r'\w+', text.lower())
    return Counter(words)
 
def get_cosine_similarity(vec1, vec2):
    intersection = set(vec1.keys()) & set(vec2.keys())
    numerator = sum([vec1[x] * vec2[x] for x in intersection])
    sum1 = sum([vec1[x]**2 for x in list(vec1.keys())])
    sum2 = sum([vec2[x]**2 for x in list(vec2.keys())])
    denominator = math.sqrt(sum1) * math.sqrt(sum2)
    return float(numerator) / denominator if denominator else 0.0
 
STOP_WORDS = {'and', 'the', 'for', 'with', 'you', 'this', 'that', 'from', 'have', 'will', 'are', 'your', 'our',
              'work', 'experience', 'looking', 'role', 'team', 'company', 'required', 'skills', 'good', 'must'}
 
def get_missing_keywords(resume_text, jd_text):
    clean_resume = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_text.lower()))
    clean_jd = set(re.findall(r'\b[a-zA-Z]{3,}\b', jd_text.lower()))
    return list((clean_jd - STOP_WORDS) - (clean_resume - STOP_WORDS))[:12]
 
def get_evidence_matches(resume_text, jd_text, limit=8):
    """For each JD keyword that IS present in the resume, find the exact resume line
    where it appears — this is the 'explainable / verbatim evidence' behind the score."""
    clean_resume = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_text.lower()))
    clean_jd = set(re.findall(r'\b[a-zA-Z]{3,}\b', jd_text.lower()))
    matched = list((clean_jd - STOP_WORDS) & (clean_resume - STOP_WORDS))
 
    resume_lines = [l.strip() for l in resume_text.split('\n') if l.strip()]
    evidence = []
    for kw in matched:
        for line in resume_lines:
            if re.search(r'\b' + re.escape(kw) + r'\b', line.lower()):
                evidence.append((kw, line))
                break
        if len(evidence) >= limit:
            break
    return evidence
 
def get_section_text(text, keywords):
    """Pulls out the block of text under a heading that matches any of `keywords`,
    stopping at the next recognizable heading."""
    text_lines = text.split('\n')
    capturing = False
    result = []
    for line in text_lines:
        l_str = line.strip().lower()
        if any(k in l_str for k in keywords) and len(l_str) < 35:
            capturing = True
            continue
        elif capturing and any(h in l_str for h in ['education', 'skills', 'experience', 'projects', 'certifications', 'summary', 'objective']) and len(l_str) < 35:
            break
        if capturing:
            result.append(line)
    return '\n'.join(result).strip()
 
def extract_resume_sections(text):
    return {
        'skills': get_section_text(text, ['skills', 'technical skills', 'technologies']),
        'experience': get_section_text(text, ['experience', 'work experience', 'employment', 'internship']),
        'projects': get_section_text(text, ['projects', 'academic projects']),
        'education': get_section_text(text, ['education', 'qualification', 'academic background']),
    }
 
def calculate_section_scores(resume_text, jd_text):
    """Breaks the single match % into per-section scores so the user can see
    WHERE the match is strong or weak, instead of one opaque number."""
    sections = extract_resume_sections(resume_text)
    v2 = text_to_vector(jd_text)
    scores = {}
    for name, sec_text in sections.items():
        content = sec_text if sec_text.strip() else resume_text  # fallback if section wasn't detected
        v1 = text_to_vector(content)
        scores[name] = round(get_cosine_similarity(v1, v2) * 100, 1)
    return scores
 
def extract_job_title(jd_text):
    lines = jd_text.strip().split('\n')
    for line in lines[:5]:
        if any(term in line.lower() for term in ["title", "role", "engineer", "analyst", "developer", "designer", "manager"]):
            return re.sub(r'[^a-zA-Z0-9\s]', '', line).strip()
    return "Software Developer"
 
def format_url(url):
    url = url.strip()
    if not url:
        return ""
    if not url.startswith("http://") and not url.startswith("https://"):
        return "https://" + url
    return url
 
def format_bullet_points(text):
    if not text.strip():
        return ""
    lines = text.strip().split('\n')
    html_out = ""
    in_list = False
    
    for line in lines:
        line_str = line.strip()
        if line_str.startswith("- ") or line_str.startswith("* ") or line_str.startswith("• "):
            if not in_list:
                html_out += "<ul style='margin: 4px 0 8px 18px; padding-left: 0; list-style-type: disc;'>"
                in_list = True
            clean_line = line_str.lstrip("-*• ").strip()
            html_out += f"<li style='margin-bottom: 3px; font-size: 12.5px;'>{clean_line}</li>"
        else:
            if in_list:
                html_out += "</ul>"
                in_list = False
            html_out += f"<p style='font-size: 12.5px; margin-bottom: 4px;'>{line_str}</p>"
            
    if in_list:
        html_out += "</ul>"
        
    return html_out
 
# ---------- No-fabrication bullet rewriter (Tab 3) ----------
def rewrite_bullet_no_fabrication(bullet):
    """Rewrites a bullet with stronger verbs WITHOUT inventing metrics.
    If the original line already has a number, that exact number is reused.
    If it doesn't, no fake percentage/number is added — impact language stays qualitative."""
    bullet = bullet.strip()
    numbers = re.findall(r'\d+(?:\.\d+)?%?\+?', bullet)
    has_metric = len(numbers) > 0
    metric = numbers[0] if has_metric else None
 
    core = re.sub(r'^(i|I)\s+(made|did|created|worked on|built|helped)\s+', '', bullet, flags=re.IGNORECASE).strip()
    core = core.rstrip('.')
    if not core:
        core = bullet.rstrip('.')
 
    if has_metric:
        opt1 = f"Engineered {core}, contributing to the {metric} result achieved."
        opt2 = f"Developed and implemented {core}, directly driving the {metric} outcome."
        opt3 = f"Built {core} end-to-end, validated by the {metric} figure reported."
    else:
        opt1 = f"Engineered {core}, with a focus on reliability and maintainable design."
        opt2 = f"Developed and implemented {core}, streamlining the overall workflow."
        opt3 = f"Built {core} from the ground up, applying best practices for performance and scalability."
 
    return opt1, opt2, opt3, has_metric
 
# ---------- DOCX / PDF export (Tab 2) ----------
def build_section_list(objective, skills, experience, projects, education, certifications, languages):
    return [
        ("Career Objective", objective),
        ("Technical Skills", skills),
        ("Experience / Internship", experience),
        ("Projects", projects),
        ("Education", education),
        ("Certifications", certifications),
        ("Languages Spoken", languages),
    ]
 
def generate_docx_resume(name, title, contact_line, objective, skills, experience, projects,
                          education, certifications, languages, declaration, dec_date, loc):
    doc = docx.Document()
    doc.add_heading(name, level=0)
    doc.add_paragraph(title)
    doc.add_paragraph(contact_line)
 
    for heading_text, content in build_section_list(objective, skills, experience, projects, education, certifications, languages):
        if content.strip():
            doc.add_heading(heading_text, level=2)
            for line in content.strip().split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith(('-', '*', '•')):
                    doc.add_paragraph(line.lstrip('-*• ').strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line)
 
    if declaration.strip():
        doc.add_heading("Declaration", level=2)
        doc.add_paragraph(declaration)
        doc.add_paragraph(f"Date: {dec_date}    Location: {loc}    Signature: {name}")
 
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
 
def generate_pdf_resume(name, title, contact_line, objective, skills, experience, projects,
                         education, certifications, languages, declaration, dec_date, loc):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(0, 10, name, align='C')
    pdf.set_font("Helvetica", "", 12)
    pdf.multi_cell(0, 8, title, align='C')
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6, contact_line, align='C')
    pdf.ln(4)
 
    for heading_text, content in build_section_list(objective, skills, experience, projects, education, certifications, languages):
        if content.strip():
            pdf.set_font("Helvetica", "B", 13)
            pdf.cell(0, 8, heading_text, ln=True)
            pdf.set_font("Helvetica", "", 10)
            for line in content.strip().split('\n'):
                line = line.strip()
                if not line:
                    continue
                bullet_prefix = "• " if line.startswith(('-', '*', '•')) else ""
                clean_line = line.lstrip('-*• ').strip()
                pdf.multi_cell(0, 6, bullet_prefix + clean_line)
            pdf.ln(2)
 
    if declaration.strip():
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, "Declaration", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, declaration)
        pdf.multi_cell(0, 6, f"Date: {dec_date}    Location: {loc}    Signature: {name}")
 
    raw = pdf.output()
    if isinstance(raw, str):
        raw = raw.encode('latin-1')
    return bytes(raw)
 
# Sidebar Inputs
st.sidebar.header("📥 Upload Documents")
resume_file = st.sidebar.file_uploader("Upload Resume (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])
resume_text_input = st.sidebar.text_area("OR Paste Resume Text", value=st.session_state["built_resume_text"], height=150)
 
extracted_resume = extract_text_from_file(resume_file)
resume_text = extracted_resume if extracted_resume != "" else resume_text_input
 
if resume_file:
    if extracted_resume != "":
        st.sidebar.success(f"✅ Loaded {len(extracted_resume.split())} words from Resume!")
    else:
        st.sidebar.warning("⚠️ Could not extract text. Please paste text in box below.")
 
st.sidebar.divider()
 
# Optional JD Toggle
has_jd = st.sidebar.radio("Do you have a Job Description (JD) to match?", ["No (Check General Resume Score)", "Yes (Compare with Job Description)"])
 
job_description = ""
if has_jd == "Yes (Compare with Job Description)":
    jd_file = st.sidebar.file_uploader("Upload Job Description (.txt, .pdf, .docx)", type=["pdf", "docx", "txt"])
    jd_text_input = st.sidebar.text_area("OR Paste JD Text", height=150)
 
    extracted_jd = extract_text_from_file(jd_file)
    job_description = extracted_jd if extracted_jd != "" else jd_text_input
 
    if jd_file and extracted_jd != "":
        st.sidebar.success(f"✅ Loaded {len(extracted_jd.split())} words from JD!")
        
# General Resume Quality Calculator (When NO JD is provided)
def calculate_general_resume_score(text):
    score = 0
    feedback = []
    word_count = len(text.split())
    
    # 1. Word Count Check (Ideal: 250 - 1000 words)
    if 250 <= word_count <= 1000:
        score += 25
        feedback.append("✅ **Ideal Length:** Resume length is optimal for ATS readable scans.")
    elif word_count < 250:
        score += 10
        feedback.append("⚠️ **Too Short:** Resume might be lacking details. Aim for at least 300+ words.")
    else:
        score += 15
        feedback.append("⚠️ **Too Long:** Try to condense your resume to 1-2 pages.")
 
    # 2. Key Sections Check
    text_lower = text.lower()
    sections = ['education', 'skills', 'experience', 'projects']
    found_sections = [s for s in sections if s in text_lower]
    score += len(found_sections) * 10  # Max 40 points
    feedback.append(f"✅ **Essential Sections Found:** Found {len(found_sections)}/4 key sections ({', '.join([s.capitalize() for s in found_sections])}).")
 
    # 3. Action Verbs Check
    action_words = ['developed', 'managed', 'created', 'built', 'designed', 'implemented', 'led', 'improved', 'analyzed', 'engineered']
    found_verbs = [v for v in action_words if v in text_lower]
    if len(found_verbs) >= 4:
        score += 20
        feedback.append(f"✅ **Strong Action Verbs:** Good usage of power words like `{', '.join(found_verbs[:4])}`.")
    else:
        score += 10
        feedback.append("💡 **Action Verbs:** Consider adding more action words (e.g., *Built, Engineered, Developed*).")
 
    # 4. Contact/Links Check
    if any(k in text_lower for k in ['email', '@', 'linkedin', 'github', 'phone']):
        score += 15
        feedback.append("✅ **Contact Information:** Contact or portfolio details detected.")
    else:
        feedback.append("⚠️ **Contact Info Missing:** Ensure your email or phone number is clearly stated.")
 
    return score, feedback
 
 
# ==================== TAB 1: RESUME ANALYZER & JOBS ====================
with tab1:
    st.subheader("📊 Resume Score & Analysis")
    
    if st.button("Run Full Analysis 🚀", type="primary"):
        if resume_text.strip() == "":
            st.error("Please upload or paste your Resume first in the Sidebar!")
        else:
            # CASE 1: USER HAS A JOB DESCRIPTION (YES)
            if has_jd == "Yes (Compare with Job Description)" and job_description.strip() != "":
                v1 = text_to_vector(resume_text)
                v2 = text_to_vector(job_description)
                match_percentage = round(get_cosine_similarity(v1, v2) * 100, 2)
                missing_skills = get_missing_keywords(resume_text, job_description)
                target_role = extract_job_title(job_description)
                encoded_role = urllib.parse.quote(target_role)
 
                col1, col2 = st.columns(2)
                with col1:
                    st.metric(label="🎯 Job Match Score", value=f"{match_percentage}%")
                    if match_percentage >= 50:
                        st.success("🎯 High Match Alignment with Job Description!")
                    elif match_percentage >= 25:
                        st.warning("⚠️ Moderate Alignment. Needs improvement.")
                    else:
                        st.error("❌ Low Alignment. Add relevant missing keywords.")
                with col2:
                    st.write("**Missing Keywords from JD:**")
                    st.write(", ".join([f"`{w}`" for w in missing_skills]) if missing_skills else "None! Excellent Job.")
 
                # ---- NEW: Explainable, section-wise breakdown ----
                st.divider()
                st.markdown("### 🧩 Section-Wise Match Breakdown")
                st.caption("The overall score is one number — this shows *where* it's coming from.")
                section_scores = calculate_section_scores(resume_text, job_description)
                sc1, sc2, sc3, sc4 = st.columns(4)
                sc1.metric("Skills Match", f"{section_scores['skills']}%")
                sc2.metric("Experience Match", f"{section_scores['experience']}%")
                sc3.metric("Projects Match", f"{section_scores['projects']}%")
                sc4.metric("Education Match", f"{section_scores['education']}%")
 
                # ---- NEW: Evidence-based matched keywords (verbatim proof) ----
                st.divider()
                st.markdown("### 🔎 Explainable Evidence — Why These Keywords Matched")
                st.caption("Every matched keyword below is backed by the exact line it was found in — no black-box scoring.")
                evidence_list = get_evidence_matches(resume_text, job_description)
                if evidence_list:
                    for kw, line in evidence_list:
                        st.write(f"✅ **`{kw}`** — found in: _\"{line}\"_")
                else:
                    st.info("No direct keyword evidence found between resume and JD.")
 
                st.divider()
                st.subheader(f"💼 Live Job Openings: {target_role}")
                c1, c2, c3 = st.columns(3)
                with c1:
                    st.link_button("Apply on LinkedIn 🚀", f"https://www.linkedin.com/jobs/search/?keywords={encoded_role}", use_container_width=True)
                with c2:
                    st.link_button("Apply on Naukri 🚀", f"https://www.naukri.com/{encoded_role.replace('%20', '-')}-jobs", use_container_width=True)
                with c3:
                    st.link_button("Apply on Indeed 🚀", f"https://www.indeed.com/jobs?q={encoded_role}", use_container_width=True)
 
            # CASE 2: USER DOES NOT HAVE A JOB DESCRIPTION (NO)
            else:
                gen_score, feedback_list = calculate_general_resume_score(resume_text)
                
                st.metric(label="📈 General Resume Quality Score", value=f"{gen_score} / 100")
                
                if gen_score >= 75:
                    st.success("🌟 Excellent Resume Structure & ATS Readiness!")
                elif gen_score >= 50:
                    st.warning("⚠️ Good Resume, but can be improved further.")
                else:
                    st.error("❌ Low Score. Consider adding missing sections or content.")
 
                st.divider()
                st.markdown("### 🔍 ATS Quality Feedback Summary:")
                for item in feedback_list:
                    st.write(item)
 
# ==================== TAB 2: LIVE RESUME EDITOR & BUILDER ====================
with tab2:
    st.subheader("✏️ Resume Upload & Interactive Editor")
    st.write("Upload your resume and click the button below — all details will automatically load into the boxes below, where you can easily modify them.")
    
    # Session State Initialization for Smart Auto-Fill
    if "ed_name" not in st.session_state: st.session_state["ed_name"] = ""
    if "ed_title" not in st.session_state: st.session_state["ed_title"] = ""
    if "ed_email" not in st.session_state: st.session_state["ed_email"] = ""
    if "ed_phone" not in st.session_state: st.session_state["ed_phone"] = ""
    if "ed_loc" not in st.session_state: st.session_state["ed_loc"] = ""
    if "ed_obj" not in st.session_state: st.session_state["ed_obj"] = ""
    if "ed_skills" not in st.session_state: st.session_state["ed_skills"] = ""
    if "ed_exp" not in st.session_state: st.session_state["ed_exp"] = ""
    if "ed_proj" not in st.session_state: st.session_state["ed_proj"] = ""
    if "ed_edu" not in st.session_state: st.session_state["ed_edu"] = ""
    if "ed_cert" not in st.session_state: st.session_state["ed_cert"] = ""
 
    # 1. File Upload Section
    uploaded_edit_file = st.file_uploader("📂 Upload Resume to Edit (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"], key="editor_file")
    
    if uploaded_edit_file is not None:
        if st.button("📥 Parse & Load Resume into Editor Boxes", type="primary"):
            extracted_raw = extract_text_from_file(uploaded_edit_file)
            if extracted_raw:
                lines = [line.strip() for line in extracted_raw.split('\n') if line.strip()]
                
                # Simple Smart Parsing Logic
                # Name (Usually first non-empty line)
                if lines:
                    st.session_state["ed_name"] = lines[0]
                
                # Email Extraction Regex/Check
                emails = re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', extracted_raw)
                if emails: st.session_state["ed_email"] = emails[0]
                
                # Phone Number Extraction Regex/Check
                phones = re.findall(r'[\+\(]?[0-9][0-9\-\s\(\)]{8,}[0-9]', extracted_raw)
                if phones: st.session_state["ed_phone"] = phones[0]
 
                # Section Extraction Logic based on Headings (uses the shared get_section_text helper)
                parsed_obj = get_section_text(extracted_raw, ['objective', 'summary', 'profile'])
                parsed_skills = get_section_text(extracted_raw, ['skills', 'technical skills', 'technologies'])
                parsed_exp = get_section_text(extracted_raw, ['experience', 'work experience', 'employment', 'internship'])
                parsed_proj = get_section_text(extracted_raw, ['projects', 'academic projects'])
                parsed_edu = get_section_text(extracted_raw, ['education', 'qualification', 'academic background'])
                parsed_cert = get_section_text(extracted_raw, ['certifications', 'certificates'])
 
                if parsed_obj: st.session_state["ed_obj"] = parsed_obj
                if parsed_skills: st.session_state["ed_skills"] = parsed_skills
                if parsed_exp: st.session_state["ed_exp"] = parsed_exp
                else: st.session_state["ed_exp"] = extracted_raw  # Fallback to full raw text if sections fail
                if parsed_proj: st.session_state["ed_proj"] = parsed_proj
                if parsed_edu: st.session_state["ed_edu"] = parsed_edu
                if parsed_cert: st.session_state["ed_cert"] = parsed_cert
 
                st.success("✅ Resume details loaded & filled into editor boxes! You can edit them below.")
            else:
                st.error("Could not extract text from this file. Please make sure it's a readable PDF or DOCX file.")
 
    st.divider()
 
    theme_choice = st.radio("🎨 Choose Resume Theme:", ["Modern Blue", "Executive Gold/Black", "Minimal Dark Header"], horizontal=True)
 
    b_col1, b_col2 = st.columns([1, 1])
 
    with b_col1:
        st.markdown("### 📝 Edit Resume Sections")
        
        # Text Inputs with Auto-Filled Values from Uploaded File
        full_name = st.text_input("Full Name", value=st.session_state["ed_name"], placeholder="e.g., John Doe")
        title = st.text_input("Professional Title", value=st.session_state["ed_title"], placeholder="e.g., Software Engineer")
        email = st.text_input("Email", value=st.session_state["ed_email"], placeholder="e.g., johndoe@example.com")
        phone = st.text_input("Phone Number", value=st.session_state["ed_phone"], placeholder="e.g., +91 9876543210")
        location = st.text_input("Location", value=st.session_state["ed_loc"], placeholder="e.g., Hyderabad, India")
        linkedin = st.text_input("LinkedIn Profile URL", value="", placeholder="https://www.linkedin.com/in/yourprofile")
        github = st.text_input("GitHub Profile URL", value="", placeholder="https://github.com/yourusername")
 
        st.markdown("---")
        objective = st.text_area("Edit Objective / Summary", value=st.session_state["ed_obj"], placeholder="Career objective or summary...", height=90)
        skills = st.text_area("Edit Skills", value=st.session_state["ed_skills"], placeholder="- Python, SQL, Streamlit...", height=110)
        experience = st.text_area("Edit Work / Internship Experience", value=st.session_state["ed_exp"], placeholder="Work experience details...", height=150)
        projects = st.text_area("Edit Projects", value=st.session_state["ed_proj"], placeholder="Project details...", height=120)
        education = st.text_area("Edit Education", value=st.session_state["ed_edu"], placeholder="Degree | College Name...", height=100)
        certifications = st.text_area("Edit Certifications (Optional)", value=st.session_state["ed_cert"], placeholder="Certifications...", height=80)
        languages = st.text_area("Languages Spoken", value="- English\n- Telugu", height=70)
        declaration = st.text_area("Declaration Statement", value="I hereby declare that all information provided is accurate to the best of my knowledge.", height=70)
        dec_date = st.text_input("Date", value="", placeholder="DD/MM/YYYY")
 
        # Fallbacks for live preview display
        p_name = full_name if full_name.strip() else "Your Full Name"
        p_title = title if title.strip() else "Professional Title"
        p_email = email if email.strip() else "email@example.com"
        p_phone = phone if phone.strip() else "+91 9876543210"
        p_loc = location if location.strip() else "City, Country"
 
        # Sync edited content with Tab 1 (Analyzer)
        edited_full_text = f"{p_name}\n{p_title}\n{objective}\n{skills}\n{experience}\n{projects}\n{education}\n{certifications}"
        if st.button("🔄 Sync Edited Resume with Analyzer"):
            st.session_state["built_resume_text"] = edited_full_text
            st.success("✅ Changes synced! Go to Tab 1 to check ATS score of this updated resume.")
 
    with b_col2:
        st.markdown("### 👁️ Live Updated Preview")
        
        if theme_choice == "Modern Blue":
            primary_color, title_color, header_bg = "#1E3A8A", "#4B5563", "transparent"
        elif theme_choice == "Executive Gold/Black":
            primary_color, title_color, header_bg = "#B45309", "#1F2937", "#FFFBEB"
        else:
            primary_color, title_color, header_bg = "#0F172A", "#64748B", "#F8FAFC"
 
        linkedin_url = format_url(linkedin) if linkedin.strip() else ""
        github_url = format_url(github) if github.strip() else ""
 
        contact_items = [f"📍 {p_loc}", f"📞 {p_phone}", f"📧 {p_email}"]
        if linkedin_url:
            contact_items.append(f"<a href='{linkedin_url}' target='_blank' style='color:{primary_color}; text-decoration: underline;'>{linkedin.strip()}</a>")
        if github_url:
            contact_items.append(f"<a href='{github_url}' target='_blank' style='color:{primary_color}; text-decoration: underline;'>{github.strip()}</a>")
        
        contact_html = " &nbsp;|&nbsp; ".join(contact_items)
        contact_plain = f"{p_loc} | {p_phone} | {p_email}"
        if linkedin.strip():
            contact_plain += f" | {linkedin.strip()}"
        if github.strip():
            contact_plain += f" | {github.strip()}"
 
        obj_s = f"<h3 style='color:{primary_color};'>Career Objective</h3>{format_bullet_points(objective)}" if objective.strip() else ""
        edu_s = f"<h3 style='color:{primary_color};'>Education</h3>{format_bullet_points(education)}" if education.strip() else ""
        skl_s = f"<h3 style='color:{primary_color};'>Technical Skills</h3>{format_bullet_points(skills)}" if skills.strip() else ""
        crt_s = f"<h3 style='color:{primary_color};'>Certifications</h3>{format_bullet_points(certifications)}" if certifications.strip() else ""
        prj_s = f"<h3 style='color:{primary_color};'>Projects</h3>{format_bullet_points(projects)}" if projects.strip() else ""
        exp_s = f"<h3 style='color:{primary_color};'>Experience / Internship</h3>{format_bullet_points(experience)}" if experience.strip() else ""
        lng_s = f"<h3 style='color:{primary_color};'>Languages Spoken</h3>{format_bullet_points(languages)}" if languages.strip() else ""
        dec_s = f"<h3 style='color:{primary_color};'>Declaration</h3><p>{declaration}</p>" if declaration.strip() else ""
 
        html_resume = f"""
        <!DOCTYPE html>
        <html>
        <head>
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Georgia&family=Calibri&display=swap');
            body {{ font-family: 'Calibri', Arial, sans-serif; margin: 0; padding: 10px; color: #222; }}
            .resume-card {{ background: #fff; border: 1px solid #ddd; padding: 25px; border-radius: 8px; }}
            .header-container {{ text-align: center; margin-bottom: 12px; background: {header_bg}; padding: 12px; border-radius: 6px; }}
            h1 {{ font-family: 'Georgia', serif; color: {primary_color}; margin: 0 0 4px 0; font-size: 26px; text-transform: uppercase; letter-spacing: 1px; }}
            .title {{ font-size: 15px; font-weight: bold; color: {title_color}; margin-bottom: 6px; }}
            .contact {{ font-size: 12px; color: #4B5563; }}
            hr {{ border: 0; border-top: 2px solid {primary_color}; margin-bottom: 15px; }}
            h3 {{ font-family: 'Georgia', serif; font-size: 14px; text-transform: uppercase; letter-spacing: 0.5px; border-bottom: 1px solid #e5e7eb; padding-bottom: 3px; margin-top: 14px; margin-bottom: 6px; }}
            p {{ font-size: 12.5px; line-height: 1.4; margin: 0; white-space: pre-line; }}
            .footer-table {{ width: 100%; margin-top: 20px; font-size: 12px; color: #555; }}
            a {{ text-decoration: none; word-break: break-all; }}
        </style>
        </head>
        <body>
            <div class="resume-card">
                <div class="header-container">
                    <h1>{p_name}</h1>
                    <div class="title">{p_title}</div>
                    <div class="contact">{contact_html}</div>
                </div>
                <hr>
                
                {obj_s} {skl_s} {exp_s} {prj_s} {edu_s} {crt_s} {lng_s} {dec_s}
                
                <table class="footer-table">
                    <tr>
                        <td><strong>Location:</strong> {p_loc}</td>
                        <td style="text-align: right;"><strong>Signature:</strong> {p_name}</td>
                    </tr>
                    <tr>
                        <td><strong>Date:</strong> {dec_date}</td>
                        <td></td>
                    </tr>
                </table>
            </div>
        </body>
        </html>
        """
        st.markdown(html_resume, unsafe_allow_html=True)
        st.write("")
 
        # ---- Export row: HTML (existing) + NEW DOCX + NEW PDF ----
        exp_col1, exp_col2, exp_col3 = st.columns(3)
 
        with exp_col1:
            st.download_button(
                label="📥 Download HTML",
                data=html_resume,
                file_name=f"{p_name.replace(' ', '_')}_Resume.html",
                mime="text/html",
                type="primary",
                use_container_width=True
            )
 
        with exp_col2:
            if docx_available:
                docx_bytes = generate_docx_resume(
                    p_name, p_title, contact_plain, objective, skills, experience,
                    projects, education, certifications, languages, declaration, dec_date, p_loc
                )
                st.download_button(
                    label="📄 Download DOCX",
                    data=docx_bytes,
                    file_name=f"{p_name.replace(' ', '_')}_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.button("📄 DOCX (install python-docx)", disabled=True, use_container_width=True)
 
        with exp_col3:
            if fpdf_available:
                pdf_bytes = generate_pdf_resume(
                    p_name, p_title, contact_plain, objective, skills, experience,
                    projects, education, certifications, languages, declaration, dec_date, p_loc
                )
                st.download_button(
                    label="🧾 Download PDF",
                    data=pdf_bytes,
                    file_name=f"{p_name.replace(' ', '_')}_Resume.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            else:
                st.button("🧾 PDF (install fpdf2)", disabled=True, use_container_width=True)
 
        if not docx_available or not fpdf_available:
            st.caption("💡 Run `pip install python-docx fpdf2` to enable all export formats.")
            with st.expander("🔧 Why is export disabled? (diagnostics)"):
                st.write(f"**python-docx available:** {docx_available}")
                if docx_import_error:
                    st.code(docx_import_error, language="text")
                st.write(f"**fpdf2 available:** {fpdf_available}")
                if fpdf_import_error:
                    st.code(fpdf_import_error, language="text")
                st.write("**Python executable in use:**")
                st.code(sys.executable, language="text")
                st.write("Run this in the SAME environment the app is launched from:")
                st.code("pip show fpdf2\npip show python-docx", language="bash")
 
# ==================== TAB 3: AI BULLET REWRITER ====================
with tab3:
    st.subheader("✨ Action-Oriented Bullet Point Improver")
    st.caption("No fabrication: rewrites only reuse numbers you already wrote — they never invent new metrics.")
    raw_bullet = st.text_input("Enter a simple project line:", "I made a python app for analysis")
    if st.button("Enhance Bullet Point ✨"):
        if raw_bullet.strip():
            opt1, opt2, opt3, has_metric = rewrite_bullet_no_fabrication(raw_bullet)
 
            if not has_metric:
                st.caption("ℹ️ No number was found in your original line, so none was invented. Add a real metric (%, count, time saved) to your resume for stronger rewrites.")
 
            st.markdown("### 🌟 Suggested Action-Oriented Options:")
            st.info(f"👉 **Option 1 (Technical Focus):** '{opt1}'")
            st.success(f"👉 **Option 2 (Action Focus):** '{opt2}'")
            st.warning(f"👉 **Option 3 (Outcome Focus):** '{opt3}'")
 
# ==================== TAB 4: SMART PLACEMENT & RED FLAGS ====================
with tab4:
    st.subheader("💡 ATS Keyword Placement & Red Flag Detector")
    if resume_text.strip() != "" and job_description.strip() != "":
        missing = get_missing_keywords(resume_text, job_description)
        st.markdown("### 📍 Where to Add Missing Keywords?")
        if missing:
            for i, kw in enumerate(missing[:6]):
                st.info(f"**Keyword:** `{kw.capitalize()}` ➔ **Suggested Bullet:** *'Successfully utilized {kw} in hands-on projects to improve workflow and optimization.'* (Add to **Skills/Projects** section)")
 
        st.divider()
        st.markdown("### 🚩 Resume Weakness & Cliché Detector")
        cliches = ['hardworking', 'honest', 'team player', 'self motivated', 'go getter', 'fast learner']
        found_cliches = [c for c in cliches if c in resume_text.lower()]
        if found_cliches:
            st.warning(f"⚠️ Found Overused Buzzwords: {', '.join(found_cliches)}")
        else:
            st.success("✅ Clean Resume! No weak buzzwords detected.")
 
# ==================== TAB 5: SALARY PREDICTOR ====================
with tab5:
    st.subheader("💰 Market Experience & Salary Range Predictor")
    if job_description.strip() != "":
        target_role = extract_job_title(job_description)
        exp_match = re.search(r'(\d+)\+?\s*(years|yrs)', job_description.lower())
        exp_years = int(exp_match.group(1)) if exp_match else 2
        base_pay = 4 + (exp_years * 2.5)
        st.metric("Predicted Annual Package (India Market)", f"₹{base_pay:.1f} LPA - ₹{base_pay + 5:.1f} LPA")
 
# ==================== TAB 6: HR COLD OUTREACH ====================
with tab6:
    st.subheader("✉️ HR Cold Outreach & Referral Generator")
    st.write("Based on your resume details, professional emails and LinkedIn messages for HR or recruiters will be generated here.")
 
    # Get candidate details from active session or input
    candidate_name = p_name if 'p_name' in locals() and p_name != "Your Full Name" else "Candidate"
    candidate_skills = skills if 'skills' in locals() and skills.strip() else "Software Development, Problem Solving"
    
    st.markdown("### 🎯 Outreach Options")
    target_company = st.text_input("Enter Target Company Name", placeholder="e.g., TCS, Google, Infosys")
    job_role = st.text_input("Enter Job Role You're Applying For", placeholder="e.g., Software Engineer, Data Analyst")
 
    if st.button("🚀 Generate Outreach Messages", type="primary"):
        if not target_company or not job_role:
            st.warning("Please enter both Target Company Name and Job Role above!")
        else:
            comp_name = target_company.strip()
            role_name = job_role.strip()
 
            st.divider()
            
            # 1. Cold Email Template
            st.markdown("### 📧 1. Personalized HR Cold Email")
            email_subject = f"Application for {role_name} Position - {candidate_name}"
            email_body = f"""Dear Hiring Team at {comp_name},
 
I hope this email finds you well.
 
I am reaching out to express my strong interest in the {role_name} position at {comp_name}. With my background in {candidate_skills[:80]}..., I am confident in my ability to add immediate value to your engineering team.
 
Key highlights from my experience:
- Hands-on expertise in key industry domains and modern tech stack.
- Proven track record of delivering clean code and working on end-to-end projects.
 
I have attached my resume for your review. I would welcome the opportunity to discuss how my skill set aligns with the goals of {comp_name}.
 
Thank you for your time and consideration.
 
Best regards,
{candidate_name}
"""
            st.code(f"Subject: {email_subject}\n\n{email_body}", language="text")
 
            # 2. LinkedIn Connection Note
            st.markdown("### 💼 2. LinkedIn Recruiter Message (Under 300 Chars)")
            linkedin_msg = f"Hi, I noticed active engineering roles at {comp_name}. As a {role_name} skilled in {candidate_skills[:40]}..., I’d love to connect and learn more about opportunities on your team. Thanks! - {candidate_name}"
            st.code(linkedin_msg, language="text")
 
            # 3. Elevator Pitch (30 Sec Intro)
            st.markdown("### 🎙️ 3. Elevator Pitch (For HR Calls)")
            pitch = f"Hi, I'm {candidate_name}. I specialize in {candidate_skills[:60]}... I recently built projects focusing on core domain solutions, and I am actively looking for a {role_name} role at {comp_name} where I can deliver impactful results."
            st.info(f"**Your Pitch:** {pitch}")
 
 
# ==================== TAB 7: AI MOCK INTERVIEW ====================
with tab7:
    st.subheader("🎤 AI Mock Interview Assistant")
    st.write("Based on your resume's projects and skills, here are the most important interview questions you might be asked.")
 
    # Determine Active Resume Text
    active_resume_text = resume_text if 'resume_text' in locals() and resume_text.strip() else edited_full_text if 'edited_full_text' in locals() else ""
 
    if st.button("🎯 Generate Interview Questions from My Resume", type="primary"):
        if not active_resume_text.strip():
            st.error("Please upload or edit a resume first to generate questions!")
        else:
            st.success("✅ Custom Interview Questions Generated based on your Resume details!")
            
            st.divider()
            
            st.markdown("### 💬 1. HR & Behavioral Questions")
            st.markdown("""
            * **Q1: Tell me about yourself and your background based on your resume.**
              > *Tip:* Focus on your skills, 1-2 major projects, and why you are suitable for this role.
            * **Q2: Why do you want to join our company?**
              > *Tip:* Mention company achievements, mission, and how your career goals align with them.
            * **Q3: Describe a challenge you faced during a project and how you solved it.**
              > *Tip:* Use the **STAR Method** (Situation, Task, Action, Result).
            """)
 
            st.markdown("---")
            
            st.markdown("### 🛠️ 2. Domain & Resume Specific Technical Questions")
            
            # Simple keyword matching to render relevant tech questions
            res_lower = active_resume_text.lower()
            
            q_count = 1
            if "python" in res_lower:
                st.markdown(f"**Q{q_count}: What are lists and tuples in Python? How do you manage memory management in Python?**")
                q_count += 1
            if "sql" in res_lower or "database" in res_lower:
                st.markdown(f"**Q{q_count}: What is the difference between WHERE and HAVING clause in SQL? Explain Joins.**")
                q_count += 1
            if "project" in res_lower:
                st.markdown(f"**Q{q_count}: Explain the architecture and technical challenges of your primary project mentioned in the resume.**")
                q_count += 1
            if "javascript" in res_lower or "react" in res_lower or "web" in res_lower:
                st.markdown(f"**Q{q_count}: Explain state vs props in React or asynchronous programming in JS.**")
                q_count += 1
            
            # Fallback general technical question
            st.markdown(f"**Q{q_count}: How do you test and debug your code before submitting a project?**")
 
            st.divider()
            st.markdown("### ✍️ Practice Your Answer")
            user_ans = st.text_area("Type your answer here to practice:", placeholder="Write your response here...")
            if st.button("Submit Answer for Review"):
                if len(user_ans.split()) > 20:
                    st.success("👍 Good response! You explained with sufficient details. Make sure to keep your tone clear and confident during the live call.")
                else:
                    st.warning("⚠️ Your response seems a bit brief. Try adding specific examples or technical details using the STAR method.")
 
